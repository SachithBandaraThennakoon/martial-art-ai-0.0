from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX = Path(__file__).with_name("Combat_Cognition_Thesis_I.T.M.S.S.B.Thennakoon.docx")
APPENDIX_URL = "https://github.com/SachithBandaraThennakoon/martial-art-ai/tree/main/research/Appendix"
HEADING = "Appendix A4: System Interface and Functional Evidence"
BODY = (
    "The publication-ready supplementary appendix contains account-name-masked "
    "screenshots with descriptive filenames. They document visible controls, "
    "state transitions and diagnostic outputs in the P001 jab-only laptop-camera "
    "case. They are observational implementation evidence, not independent "
    "measurements of accuracy, effectiveness, causal reasoning or human-equivalent "
    "cognition. To avoid duplicating the same screenshots inside the report, the "
    "curated images are provided through the repository link below."
)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


document = Document(DOCX)

# Remove the twelve externalized A9 entries from the in-report List of Figures.
for paragraph in list(document.paragraphs):
    if paragraph.text.strip().startswith("Figure A9."):
        remove_paragraph(paragraph)

paragraphs = document.paragraphs
heading_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == HEADING)
intro = paragraphs[heading_index + 1]

# A9 is the final report section. Remove its duplicated screenshot pages and captions.
body = document._element.body
intro_element = intro._element
remove_after = False
for child in list(body):
    if remove_after:
        body.remove(child)
    elif child is intro_element:
        remove_after = True

intro.text = BODY
link_paragraph = document.add_paragraph()
link_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
link_paragraph.add_run("Supplementary appendices: ").bold = True
add_hyperlink(link_paragraph, APPENDIX_URL, APPENDIX_URL)

document.save(DOCX)
print(DOCX)
