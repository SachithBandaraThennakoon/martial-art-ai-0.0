from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[4]
OUT = ROOT / "research/outputs/thesis/20260801"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Combat_Cognition_Thesis_I.T.M.S.S.B.Thennakoon.docx"

TITLE = "COMBAT COGNITION: A HYBRID COMPUTATIONAL FRAMEWORK TOWARD MARTIAL-ARTIST COGNITIVE SIMULATION"
AUTHOR = "I.T.M.S.S.B. THENNAKOON"
REG_NO = "DTS2401"
DATE = "1 AUGUST 2026"
APPENDIX_BASE_URL = "https://github.com/SachithBandaraThennakoon/martial-art-ai/tree/main/research/Appendix"
APPENDIX_FOLDERS = {
    "Appendix A1: Architecture and Implementation Evidence": "A1-Architecture-and-Implementation",
    "Appendix A2: ACP-STGAT Reproducibility Record": "A2-ACP-STGAT-Reproducibility",
    "Appendix A3: Phase-Classifier Reproducibility Record": "A3-Phase-Classifier-Reproducibility",
    "Appendix A4: System Interface and Functional Evidence": "A4-System-Interface-and-Functional-Evidence",
    "Appendix A5: P001 Protocol and Curated Evidence": "A5-P001-Protocol-and-Curated-Evidence",
    "Appendix A6: Software Verification": "A6-Software-Verification",
    "Appendix A7: Database Export and Failure Evidence": "A7-Database-Export-and-Failure-Evidence",
    "Appendix A8: Literature and Claim Audit": "A8-Literature-and-Claim-Audit",
    "Appendix A9: Evaluation Actions and Evidence Locks": "A9-Evaluation-Actions-and-Evidence-Locks",
}


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_fixed(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(round(w * 1440) for w in widths_inches)))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[i])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(round(widths_inches[i] * 1440)))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, title, headers, rows, widths, note=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(title), 12, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9E2F3")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        set_font(para.add_run(str(header)), 9, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            if i > 0 and isinstance(value, (int, float)):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(para.add_run(str(value)), 9)
    set_table_fixed(table, widths)
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(f"Note: {note}"), 9, italic=True)
    return table


def add_body(doc, text, citation=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    set_font(p.add_run(text), 12)
    if citation:
        set_font(p.add_run(f" {citation}"), 12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    set_font(p.add_run(text), 12)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    set_font(p.add_run(text), 12)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        set_font(r, 12, bold=True, italic=(level == 3))
    if not p.runs:
        set_font(p.add_run(text), 12, bold=True, italic=(level == 3))
    else:
        p.runs[0].text = text
    return p


def add_chapter(doc, number, title, force_page_break=True):
    if force_page_break and len(doc.paragraphs) > 0:
        doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f"CHAPTER {number}\n{title.upper()}")
    set_font(run, 14, bold=True)
    return p


def add_figure(doc, path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.style = "Caption"
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.keep_with_next = False
    c.paragraph_format.space_before = Pt(3)
    c.paragraph_format.space_after = Pt(8)
    set_font(c.add_run(caption), 10, italic=True)


def add_field(paragraph, instruction, placeholder):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, 12)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def page_number_footer(section, roman=False):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE", "1")
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), "lowerRoman" if roman else "decimal")
    pg_num.set(qn("w:start"), "1")


def centered_page_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_font(p.add_run(text.upper()), 14, bold=True)


def add_index_line(doc, label, page):
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(5)
    p.paragraph_format.line_spacing=1.0
    tabs=p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(5.85))
    set_font(p.add_run(f"{label}\t{page}"),11)


doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.left_margin = Mm(40)
sec.right_margin = Mm(15)
sec.top_margin = Mm(25)
sec.bottom_margin = Mm(25)
sec.header_distance = Mm(12)
sec.footer_distance = Mm(12)

# University-thesis override of the narrative preset: exact institutional page,
# margin, Times New Roman, heading and spacing rules take precedence.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)
for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12), ("Caption", 10)):
    st = styles[name]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.font.size = Pt(size)
    st.font.bold = name != "Caption"
    st.font.color.rgb = RGBColor(0, 0, 0)

# Title page: course-specific DS5299 layout.
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(34); set_font(p.add_run(TITLE), 14, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(22); set_font(p.add_run("INDEPENDENT STUDY REPORT SUBMITTED BY"), 12, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(6); set_font(p.add_run(AUTHOR), 12, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(38); set_font(p.add_run(f"REGISTRATION NUMBER: {REG_NO}"), 12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8); set_font(p.add_run("TO THE BOARD OF STUDY IN STATISTICS AND COMPUTER SCIENCE"), 12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(30); set_font(p.add_run("POSTGRADUATE INSTITUTE OF SCIENCE"), 12, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(46); set_font(p.add_run("DS 5299 - INDEPENDENT STUDY"), 12, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run("UNIVERSITY OF PERADENIYA\nSRI LANKA\n" + DATE), 12, bold=True)

doc.add_page_break(); centered_page_heading(doc, "Declaration")
add_body(doc, "I hereby declare that the work reported in this independent study report was carried out by me and that the results are my own independent research, except where due acknowledgement is made in the text. This report has not been submitted, either in whole or in part, for any other degree or qualification.")
add_body(doc, "Candidate: I.T.M.S.S.B. Thennakoon\nRegistration number: DTS2401\nSignature: ______________________________\nDate: ______________________________")
add_body(doc, "Supervisor certification: ______________________________\nSignature: ______________________________\nDate: ______________________________")

doc.add_page_break(); centered_page_heading(doc, "Abstract")
abstract = (
    "Combat Cognition is a hybrid computational framework intended to connect markerless human-pose perception, temporal movement interpretation, short-term motion prediction, biomechanical rules, multi-timescale state, situation-aware evidence gating, memory and coaching feedback. This independent study designed, implemented and evaluated the framework as a bounded step toward martial-artist cognitive simulation. The implemented artifact uses a laptop camera and MediaPipe-derived landmarks, an ACP-STGAT future-pose model, a temporal phase-classification and duration-aware decoding pipeline, deterministic biomechanics and coaching rules, four levels of contextual state, and persistent session analytics. Evidence was deliberately separated into four tiers. On a participant-held-out Human3.6M-17 protocol, ACP-STGAT achieved a mean normalized MPJPE/ADE of 0.07839 and final displacement error of 0.13639 across three seeds, outperforming last-pose and constant-velocity baselines. The phase classifier achieved 87.84% mean accuracy, 90.84% balanced accuracy and 0.8915 macro F1 on generated bootstrap sessions; these values validate generator-defined pipeline behaviour rather than real-human accuracy. Automated verification passed 129 of 129 frontend assertions and 24 of 24 backend tests. A single-participant expert feasibility/self-evaluation used a confirmed jab session captured by the researcher, developer and martial-arts expert. The selected practice-42 session contained three confirmed movement clusters, 250 frames, 8.3 seconds of capture, 96% recorded tracking quality and 22 frames changed by post-session decoding. The same evidence exposed an inconsistency between canonical, rule-engine and database repetition representations, retained as a failure mode. The study supports an implemented, evidence-bounded prototype architecture, not effectiveness, independent usability validation, universal technique correctness, calibrated three-dimensional biomechanics, operational large-language-model reasoning or human-equivalent cognition. Future evaluation requires independent annotation, grouped real jab data, model ablation, end-to-end timing, condition-labelled comparison and broader participants."
)
add_body(doc, abstract)

doc.add_page_break(); centered_page_heading(doc, "Acknowledgements")
add_body(doc, "I acknowledge the academic guidance and institutional support associated with DS 5299 - Independent Study at the Postgraduate Institute of Science, University of Peradeniya. I also acknowledge the researchers whose published work provided the foundations for markerless pose estimation, skeleton graph learning, motion prediction, temporal segmentation, situation awareness, design-science evaluation and sports biomechanics used in this study. The practical system reflects sustained study and martial-arts experience; responsibility for the design decisions, implementation, evaluation boundaries and interpretation remains with the author.")

doc.add_page_break(); centered_page_heading(doc, "Table of Contents")
for label,page in [
    ("Declaration","ii"),("Abstract","iii"),("Acknowledgements","iv"),("List of Figures","vi"),("List of Tables","vii"),("List of Abbreviations","viii"),
    ("Chapter 1 - Introduction","1"),("Chapter 2 - Literature Review","4"),("Chapter 3 - Methodology","9"),("Chapter 4 - Results and Discussion","14"),("Chapter 5 - Conclusions and Recommendations","22"),("References","25"),("Appendices","27")
]: add_index_line(doc,label,page)
doc.add_page_break(); centered_page_heading(doc, "List of Figures")
for text,page in [
    ("Figure 3.1 Combat Cognition implemented architecture and evidence flow","10"),
    ("Figure 4.1 ACP-STGAT normalized prediction error by forecast horizon","14"),
    ("Figure 4.2 Generated-bootstrap temporal phase confusion matrix","16"),
    ("Figure 4.3 P001 practice-42 post-session three-cluster timeline","18"),
]: add_index_line(doc,text,page)
doc.add_page_break(); centered_page_heading(doc, "List of Tables")
for text,page in [
    ("Table 1.1 Research objectives and evidence boundaries","2"),
    ("Table 2.1 Literature-to-framework synthesis","7"),
    ("Table 3.1 Combat Cognition component architecture","10"),
    ("Table 3.2 Evaluation evidence hierarchy","12"),
    ("Table 4.1 ACP-STGAT model results and baselines","14"),
    ("Table 4.2 Temporal phase-classifier generated-bootstrap results","15"),
    ("Table 4.3 Automated software verification","17"),
    ("Table 4.4 P001 practice-42 evidence summary","17"),
    ("Table 4.5 Retained data-layer inconsistency","18"),
    ("Table 4.6 Controlled claim readiness","19"),
]: add_index_line(doc,text,page)
doc.add_page_break(); centered_page_heading(doc, "List of Abbreviations")
add_table(doc, "Abbreviations", ["Abbreviation", "Meaning"], [
    ["ACP-STGAT", "Action-Conditioned Physics-Informed Spatio-Temporal Graph Attention Transformer"],
    ["ADE", "Average displacement error"], ["FDE", "Final displacement error"],
    ["F1", "Harmonic mean of precision and recall"], ["LLM", "Large language model"],
    ["MPJPE", "Mean per-joint position error"], ["ONNX", "Open Neural Network Exchange"],
    ["P001", "Single researcher/developer/martial-arts expert participant"],
    ["SA", "Situation awareness"], ["ST-GCN", "Spatial-temporal graph convolutional network"],
    ["TCN", "Temporal convolutional network"],
], [1.45, 4.65])
page_number_footer(sec, roman=True)

# Main section with Arabic numbering.
main = doc.add_section(WD_SECTION.NEW_PAGE)
main.page_width = Mm(210); main.page_height = Mm(297); main.left_margin = Mm(40); main.right_margin = Mm(15); main.top_margin = Mm(25); main.bottom_margin = Mm(25); main.header_distance = Mm(12); main.footer_distance = Mm(12)
main.footer.is_linked_to_previous = False
page_number_footer(main, roman=False)

add_chapter(doc, 1, "Introduction", force_page_break=False)
add_heading(doc, "1.1 Background", 2)
add_body(doc, "Digital movement analysis has progressed from isolated image classification toward systems that track landmarks, represent temporal structure and support interactive feedback. In a martial-arts setting, however, a visible body configuration is rarely meaningful by itself. The same elbow angle can indicate preparation, extension, impact proximity or recovery depending on the surrounding movement. Coaching decisions also depend on tracking reliability, the ordering of actions, previous repetitions and the learner's recent history. These requirements motivate an integrated architecture rather than a single pose classifier.")
add_body(doc, "Combat Cognition was conceived as a computational framework that links perception, comprehension and limited projection across several timescales. It receives monocular laptop-camera input, extracts body and hand evidence, evaluates expert-authored biomechanical rules, estimates temporal phase, forecasts near-future skeletal motion, maintains action/session/user state, selects an evidence-ranked correction and stores the resulting session history. The term cognition is used in a bounded computational sense. The artifact implements selected information-processing functions associated with interpreting movement and choosing feedback; it does not reproduce consciousness, biological cognition or the complete tactical capability of a human martial artist.")
add_body(doc, "The jab was selected as the single evaluation technique. It offers a compact but temporally ordered movement involving preparation, entry, extension, peak, retraction and recovery, while still requiring coordination of the shoulder, elbow, guard and body position. The framework is intended to be extensible, but evidence collected in this study applies only to the jab case, one researcher-expert participant and one laptop-camera configuration.")

add_heading(doc, "1.2 Research Problem", 2)
add_body(doc, "Many pose and exercise-feedback systems focus on detecting landmarks, classifying an action or checking a small set of angles. These are valuable capabilities, but they do not by themselves connect short-term anticipation, ordered phase interpretation, evidence confidence, multi-repetition context and longer-term learner memory. A system that reacts to each frame independently can issue contradictory or mistimed corrections. Conversely, a learned model without deterministic evidence gates may produce a plausible prediction that should not be trusted under occlusion, tracking loss or domain shift.")
add_body(doc, "The research problem is therefore how to design and evaluate a hybrid perception-awareness-decision framework that integrates learned temporal evidence with deterministic constraints while remaining explainable about what has and has not been validated. The problem includes a methodological challenge: model benchmarks, software tests and a single expert case measure different things. Treating them as one overall accuracy would create a scientifically invalid result.")

add_heading(doc, "1.3 Aim", 2)
add_body(doc, "The aim of this study is to design, implement and evaluate the Combat Cognition Framework as a core computational architecture for perceiving human movement, developing bounded situation awareness, assessing a martial-arts technique, making coaching decisions and providing personalized feedback.")

add_heading(doc, "1.4 Research Question", 2)
add_body(doc, "How can a hybrid perception-awareness-decision framework simulate selected computational functions of a martial artist when interpreting and assessing a martial-arts technique?")

add_heading(doc, "1.5 Objectives", 2)
objectives = [
    ["O1", "Implement laptop-camera body/hand perception, visibility-aware motion context and expert-authored jab-form rules.", "Implementation feasibility; not independent landmark or form accuracy."],
    ["O2", "Implement ACP-STGAT and evaluate a 30-frame forecast using participant-held-out benchmark data, simple baselines, robustness and ONNX checks.", "Normalized Human3.6M-17 evidence; not live jab prediction."],
    ["O3", "Implement ordered temporal phase classification and decoding and verify the pipeline using generated-bootstrap metrics and P001 post-session observations.", "Generated structure and self-observation; not real annotated accuracy."],
    ["O4", "Integrate L1-L4 computational context, evidence-gated situation awareness, memory and a structured coaching context.", "Computational architecture; not human awareness."],
    ["O5", "Implement deterministic rule/template coaching through voice/UI pathways with persistent session analytics.", "No operational OpenAI LLM and no independent feedback-correctness result."],
    ["O6", "Evaluate the artifact using separated model, software, framework-case and failure-mode evidence.", "P001 jab-only feasibility; not effectiveness or combined accuracy."],
]
add_table(doc, "Table 1.1 Research objectives and evidence boundaries", ["ID", "Objective", "Boundary"], objectives, [0.45, 3.6, 2.05])

add_heading(doc, "1.6 Scope and Contributions", 2)
add_body(doc, "The empirical scope is intentionally narrow: participant P001 is the researcher, developer and martial-arts expert; the technique is a jab; and sensing uses a laptop camera. The work is classified as a single-participant expert feasibility/self-evaluation case embedded in a design-science study. Original recordings remain private and are excluded from the submitted report. Screenshots and derived session evidence are restricted to functional illustration.")
add_body(doc, "The study contributes an implemented multi-layer architecture, two deployed temporal models, explicit evidence gates, deterministic coaching and a reproducible claim-control protocol. It also contributes negative and limiting evidence: generated phase data are not presented as real accuracy; model-only latency is not end-to-end latency; a database integration inconsistency is retained; and the absence of an operational LLM is reported rather than obscured.")

add_heading(doc, "1.7 Structure of the Report", 2)
add_body(doc, "Chapter 2 reviews the evidence that motivates markerless perception, skeleton graph modelling, future-motion prediction, temporal segmentation, monocular uncertainty, situation awareness, design-science evaluation and jab biomechanics. Chapter 3 describes the research design and implemented architecture. Chapter 4 reports offline model evidence, software verification, the P001 framework case, failure modes and evidence-to-claim reconciliation. Chapter 5 answers the research question, summarizes contributions and identifies future evaluation.")

add_chapter(doc, 2, "Literature Review")
add_heading(doc, "2.1 Markerless Perception and On-Device Pose Estimation", 2)
add_body(doc, "A perception layer is required before higher-order temporal reasoning can occur. BlazePose demonstrates a lightweight detector-tracker pipeline that estimates 33 body keypoints for a single person and is designed for real-time use (Bazarevsky et al., 2020). BlazePose GHUM extends this direction toward body and hand landmarks from monocular RGB input and emphasizes on-device operation (Grishchenko et al., 2022). MediaPipe Hands provides a related 21-landmark hand model suitable for guard and hand-readiness diagnostics (Zhang et al., 2020). MediaPipe's graph-based framework further supports modular perception pipelines in which sensing, transformation and inference components are combined as processing graphs (Lugaresi et al., 2019).")
add_body(doc, "These sources justify the technical feasibility of laptop-camera landmark extraction, but they do not validate the downstream Combat Cognition rules. Vendor or benchmark keypoint measures are not jab-form accuracy. Hand landmarks may be occluded, depth is uncertain, and a single-person tracker can fail under fast motion or unusual framing. The local system therefore retains visibility, tracking-quality and failure evidence instead of assuming every landmark is correct.")

add_heading(doc, "2.2 Skeletons as Spatial-Temporal Graphs", 2)
add_body(doc, "Skeleton sequences have a natural graph interpretation: joints are nodes, anatomical links form spatial edges and repeated joints across frames form temporal edges. ST-GCN established this representation for action recognition (Yan, Xiong and Lin, 2018). Two-stream adaptive graph convolution showed that learned topology and bone vectors can complement fixed physical structure (Shi et al., 2019), while multi-scale spatial-temporal graph convolution motivated relationships beyond immediate neighbours (Chen et al., 2021). Actional-structural GCNs similarly distinguish structural and task-dependent connections and use future-pose prediction as an auxiliary representation objective (Li et al., 2019).")
add_body(doc, "The literature supports graph-based representation as a design precedent, not direct validation of ACP-STGAT. Most cited studies address recognition on large skeleton datasets rather than regression of future coordinates in a live jab. Learned edges must also be treated as computational dependencies rather than causal biomechanical pathways. The project accordingly evaluates the predictor on its own stated protocol and avoids transferring external accuracy figures.")

add_heading(doc, "2.3 Future-Motion Prediction", 2)
add_body(doc, "Future-motion prediction seeks a sequence of poses beyond the observed window. Martinez, Black and Romero (2017) demonstrated that a simple last-pose or zero-velocity predictor can be difficult to beat at short horizons, making simple baselines essential. Mao et al. (2019) modelled joint trajectories using temporal transformations and learned graph dependencies and reported error across forecast horizons. These findings support the use of horizon-specific displacement errors and baseline comparisons rather than an invented accuracy percentage.")
add_body(doc, "Combat Cognition's ACP-STGAT is evaluated using normalized mean per-joint displacement, final displacement and bone-length consistency. Because coordinates are normalized, errors cannot be labelled millimetres or centimetres. Human3.6M provides calibrated multi-view motion-capture evidence for general human activities (Ionescu et al., 2014), but its controlled activities and 17-joint representation differ from a MediaPipe-33 laptop-camera jab. The benchmark can test forecasting capability while leaving a clear domain gap.")

add_heading(doc, "2.4 Temporal Phase Classification and Segmentation", 2)
add_body(doc, "A martial-arts technique is temporally ordered, and frame accuracy alone can conceal boundary jitter, over-segmentation or impossible transitions. MS-TCN uses multi-stage temporal convolution and a smoothing objective to refine framewise action segments (Abu Farha and Gall, 2019). ASFormer combines local temporal priors, hierarchical attention and iterative decoding and also illustrates the difficulty of high-capacity temporal models when labelled data are limited (Yi, Wen and Jiang, 2021).")
add_body(doc, "These studies motivate reporting balanced accuracy, macro and per-class F1, confusion, boundary timing and repetition recovery. They do not make generated bootstrap data equivalent to human annotation. The current phase classifier is therefore presented as a complete synthetic pipeline evaluation. Its strong boundary scores describe generator-defined, smoothly separated phases; human movement may contain gradual, ambiguous and participant-dependent transitions.")

add_heading(doc, "2.5 Monocular Three-Dimensional Uncertainty", 2)
add_body(doc, "Monocular images do not uniquely determine three-dimensional pose. Controlled datasets such as Human3.6M use multiple cameras and motion capture (Ionescu et al., 2014), whereas MPI-INF-3DHP broadens appearance and environment variation but remains a curated benchmark with limited actors (Mehta et al., 2017). The 3DPW dataset combines video with body-worn inertial units to strengthen in-the-wild ground truth (von Marcard et al., 2018). Probabilistic monocular pose research explicitly represents multiple feasible three-dimensional hypotheses for the same two-dimensional observation (Wehrbein et al., 2021).")
add_body(doc, "This literature establishes a mandatory measurement boundary. MediaPipe-relative coordinates and camera-derived angles are useful computational features, but they are not independently calibrated biomechanics. Depth ambiguity, clothing, viewpoint, self-occlusion and high-speed movement affect confidence. Physical-unit, force, injury or energy-expenditure claims would require calibrated external sensors and a different protocol.")

add_heading(doc, "2.6 Situation Awareness and Multi-Timescale Context", 2)
add_body(doc, "Endsley's situation-awareness model distinguishes perception of elements, comprehension of their meaning and projection of likely future status (Endsley, 1995). This provides a useful conceptual lens for the framework: landmarks and motion contribute perception; phase, form and session state contribute computational interpretation; and ACP-STGAT contributes a bounded projection signal. The mapping does not imply that the software possesses human awareness. Goals, expertise, attention and mental models are properties of human cognition that are not reproduced merely by naming software layers after them.")
add_body(doc, "Combat Cognition adds multiple time horizons: L1 represents immediate motion, L2 represents action and projection, L3 represents the session and repetitions, and L4 represents stored learner history. This hierarchy addresses a gap between systems that end at action classification and coaching systems that must remember repeated errors or adjust pacing. Longitudinal validity remains untested in the present single-case evaluation.")

add_heading(doc, "2.7 Design-Science Evaluation and Reflexivity", 2)
add_body(doc, "Design-science research evaluates purposeful artifacts through a combination of relevance, rigor, construction and evaluation (Hevner et al., 2004). This orientation fits Combat Cognition because the contribution is an integrated artifact and associated design knowledge, not a population-level clinical intervention. Software tests, model experiments, an expert case and failure evidence can jointly characterize the artifact, provided they are not collapsed into a single metric.")
add_body(doc, "The researcher also supplied the martial-arts expertise used to formulate phase ordering, feedback priorities and thresholds. This creates valuable domain knowledge but also reflexive bias: the same individual designed, implemented and evaluated the system. The study therefore uses the term expert self-evaluation, identifies thresholds as expert-authored and treats independent expert triangulation as future work.")

add_heading(doc, "2.8 Jab Biomechanics and Context", 2)
add_body(doc, "Punch kinematics depend on punch type and execution context. Piorkowski, Lees and Barton (2011) reported differences between single and combination punches, cautioning against transferring one protocol to another. Lockwood and Tant (1997) examined jab mechanics using two-dimensional kinematics, electromyography and an instrumented bag in a small sample, supporting a coordinated multi-joint view while also illustrating sampling and measurement limitations. Ald'ily et al. (2023) analysed straight-punch variants with consumer video and considered shoulder, elbow, pelvis and lower-limb variables, but their values cannot define universal jab thresholds.")
add_body(doc, "The literature supports including multiple joints and contextual variables, not treating elbow extension as the whole technique. Combat Cognition accordingly evaluates elbow, shoulder and guard-related evidence within an ordered sequence. Numerical coaching ranges remain prototype expert-authored tolerances. The camera cannot measure punch force, muscle activation, metabolic cost or injury prevention.")

add_heading(doc, "2.9 Literature Synthesis and Research Gap", 2)
lit_rows = [
    ["Markerless pose", "Real-time body/hand landmarks", "Task and depth uncertainty", "Perception with confidence/visibility"],
    ["Skeleton graphs", "Spatial-temporal joint representation", "Recognition differs from prediction", "ACP/phase design rationale"],
    ["Motion prediction", "Baselines and horizon errors", "Benchmark/domain dependence", "ACP evaluation protocol"],
    ["Temporal segmentation", "Sequence and boundary metrics", "Requires labelled real sequences", "Phase pipeline and decoder"],
    ["Situation awareness", "Perception-comprehension-projection lens", "Human theory does not prove machine awareness", "Bounded L1-L4 evidence gates"],
    ["Jab biomechanics", "Multi-joint, context-dependent movement", "Small/heterogeneous studies", "Variable selection, not universal thresholds"],
]
add_table(doc, "Table 2.1 Literature-to-framework synthesis", ["Area", "Supported proposition", "Key limitation", "Framework use"], lit_rows, [1.15, 1.75, 1.55, 1.65])
add_body(doc, "The synthesis identifies an integration gap rather than claiming that no related systems exist. The literature contains mature work on pose estimation, graph action recognition, motion prediction and temporal segmentation, but each external result has a different task, dataset and measurement boundary. Combat Cognition's contribution is the explicit integration of those computational functions with deterministic biomechanics, multi-timescale state, evidence gating, session memory and coaching. The empirical contribution is feasibility evidence for that integration, not proof that it outperforms all alternative coaching systems.")
add_body(doc, "The ICCV paper named STGAT predicts pedestrian trajectories rather than articulated skeletal pose (Huang et al., 2019). It is retained as a terminology exclusion because acronym similarity can be misleading. Its displacement results are not used to validate ACP-STGAT.")

add_chapter(doc, 3, "Methodology")
add_heading(doc, "3.1 Research Design", 2)
add_body(doc, "The study follows an expert-informed design-science case approach. The artifact was designed from literature, software requirements and the researcher's martial-arts knowledge; implemented as an interactive laptop-camera application; and evaluated through separate evidence tiers. The empirical case is P001 only. P001 is simultaneously the researcher, developer and martial-arts expert, so observations are treated as reflective expert self-evaluation rather than independent participant evidence.")
add_body(doc, "The evaluation was intentionally constrained after practical limitations made a multi-participant study infeasible. The protocol was not converted into an effectiveness study. Instead, the final question became whether the system components functioned, whether model pipelines produced reproducible measured outputs under their declared datasets, whether an end-to-end jab session could be demonstrated, and whether failures and unsupported claims could be identified.")

add_heading(doc, "3.2 Framework Architecture", 2)
add_figure(doc, ROOT / "research/figures/verified/20260801/F3_combat_cognition_architecture_evidence_flow.png", "Figure 3.1 Combat Cognition implemented architecture and evidence flow. Solid boxes are implemented; the dashed LLM component is planned and not operational. Evidence classes remain separate.", 6.05)
add_body(doc, "Figure 3.1 shows the implemented data path. Laptop-camera frames feed MediaPipe-derived perception. L1 calculates immediate motion and deterministic biomechanical evidence. L2 combines phase estimation, ordered decoding, future-pose prediction and forecast trust. L3 aggregates repetition and session trends. L4 stores longer-term learner context. A situation-awareness layer selects an evidence-supported target, a structured context packet crosses the coaching boundary, and deterministic rules/templates produce voice and interface feedback. Timeline and analysis services persist session evidence.")

components = [
    ["Perception", "MediaPipe pose/hand signals", "33 body landmarks, readiness proxies, visibility", "Implemented; reliability unquantified"],
    ["Biomechanics", "Angle and rule engine", "Joint angles, issues, form evidence", "Expert-authored thresholds"],
    ["L1", "Immediate motion layer", "Recent motion and tracking context", "Camera dependent"],
    ["L2", "ACP-STGAT and phase/decoder", "Future pose, phase, risk and trusted forecast", "Benchmark/generated evidence"],
    ["L3", "Session layer", "Repetition, trend, consistency, fatigue proxy", "Not physiological fatigue"],
    ["L4", "User memory", "Weakness, mastery and progression structure", "Longitudinal validity absent"],
    ["Situation", "Evidence gates", "Priority target, decision and next action", "No human-awareness claim"],
    ["Reasoning", "Rule/template coach", "Voice/UI correction", "No operational LLM"],
]
add_table(doc, "Table 3.1 Combat Cognition component architecture", ["Layer", "Implementation", "Output", "Boundary"], components, [0.75, 1.65, 2.05, 1.65])

add_heading(doc, "3.3 Perception and Biomechanical Processing", 2)
add_body(doc, "The perception pipeline processes a single person in browser video and extracts body landmarks, visibility and derived hand/face readiness features. Landmark histories are normalized for the downstream temporal components. The angle engine uses joint triplets to calculate form evidence, while a deterministic rule evaluator compares evidence against technique targets. Tracking confidence and short-gap handling are retained because a geometrically plausible angle can still be unreliable when one landmark is occluded.")
add_body(doc, "The rule system prioritizes contextual correction. For example, an elbow-extension issue is not evaluated independently of shoulder and guard evidence. During the documented Train Hard interaction, the system waited for the user's response, accepted voice input, prioritized the left elbow, reported angle-related feedback and progressed after detecting a movement transition rather than merely asking for a larger angle. Shoulder reduction was interpreted in relation to body guard. These observations demonstrate the intended interaction logic but do not constitute independent coaching accuracy.")

add_heading(doc, "3.4 ACP-STGAT Future-Pose Model", 2)
add_body(doc, "ACP-STGAT consumes a recent 17-joint three-dimensional skeleton window and predicts 30 future frames. The evaluation used a participant-held-out Human3.6M protocol with up to 10,000 test windows and three training seeds (42, 43 and 44). Performance was measured through normalized MPJPE/ADE, final displacement error and bone-length mean absolute error. Last-pose and constant-velocity predictors were included as required simple baselines. Robustness trials added coordinate noise and simulated missing landmarks. The selected model was exported to ONNX and compared numerically with its PyTorch output.")
add_body(doc, "The live application uses MediaPipe-33 landmarks, whereas the offline benchmark uses a 17-joint representation. The benchmark therefore evaluates general forecasting capability and export readiness, not live jab prediction. No error is converted to a percentage or physical unit. The model-only latency excludes camera capture, MediaPipe processing, visualization and coaching logic.")

add_heading(doc, "3.5 Temporal Phase Classification and Ordered Decoding", 2)
add_body(doc, "The temporal classifier combines spatial skeleton processing with temporal convolution to produce probabilities for tracking lost, preparation, entry, execution, peak, retraction and recovery. A duration-aware decoder and application state machine constrain invalid transitions and support repetition completion. The training/evaluation bootstrap contained 48 generated sessions: 24 jab and 24 front-kick sessions. Three random seeds were evaluated, and deployment selection used validation macro F1 rather than test performance.")
add_body(doc, "Evaluation included accuracy, balanced accuracy, macro precision/recall/F1, weighted F1, per-phase metrics, confusion, boundary precision/recall/timing, coarse repetition diagnostics, corruption robustness and ONNX parity/latency. Because generated sessions share a procedural generator family, the evaluation tests pipeline consistency rather than real-world generalization. P001 post-session decoding is used only as qualitative functional evidence.")

add_heading(doc, "3.6 Multi-Level Context and Situation-Awareness Gates", 2)
add_body(doc, "L1 maintains current and recent motion evidence. L2 combines current action state, phase probabilities, biomechanical issues and forecast confidence. L3 aggregates repetitions, repeated issues, consistency and a non-medical fatigue-risk proxy. L4 stores learner history and mastery-oriented state. The situation layer receives L1-L4 evidence, determines the current target and suppresses feedback that lacks sufficient support. This structure operationalizes perception, computational comprehension and bounded projection while keeping deterministic gates authoritative.")

add_heading(doc, "3.7 Coaching, Voice Interaction and Memory", 2)
add_body(doc, "The operational coaching layer is deterministic. A structured context packet summarizes relevant L1-L4 evidence and selected actions; backend rule/template generators produce the final coaching message. The interface supports voice/microphone responses, waits for the user where required and allows explicit commands such as moving to the next step or training again. No operational OpenAI API call or deployed LLM was found in the audited repository. A replaceable LLM remains a future component behind the same structured boundary and safety gates.")
add_body(doc, "Session recording stores timelines, repetitions, feedback events, analytics and selected landmark tapes. The dashboard and analysis interfaces expose filters, movement clusters, selected-frame evidence, diagnostics and multi-level skeleton visualizations. Stored history supports later review, but development records are not automatically treated as controlled research data.")

add_heading(doc, "3.8 Evidence Hierarchy and Evaluation Protocol", 2)
hierarchy = [
    ["Offline model", "ACP benchmark; generated phase bootstrap", "Forecasting and pipeline behaviour under declared datasets", "Live-system or real-human accuracy"],
    ["Software", "Frontend/backend tests and implementation audit", "Verified code paths and contracts", "Task accuracy or effectiveness"],
    ["Framework case", "P001 practice-42, screenshots and curated export", "Single expert feasibility/self-evaluation", "Independent usability or generalization"],
    ["Literature", "22-source matrix and 39 extracts", "Design rationale and claim boundaries", "Local empirical performance"],
]
add_table(doc, "Table 3.2 Evaluation evidence hierarchy", ["Tier", "Evidence", "Supports", "Does not support"], hierarchy, [1.0, 1.7, 1.8, 1.6])
add_body(doc, "The protocol prohibits combining these tiers into one overall accuracy. Each metric is reported with its dataset, unit and evidence status. Historical sessions lacking an experiment condition are retained for availability and failure analysis but excluded from the headline case. Missing measures, such as end-to-end response time and structured usability ratings, are not reconstructed retrospectively.")

add_heading(doc, "3.9 Privacy, Ethics and Reflexivity", 2)
add_body(doc, "The participant identifier P001 replaces account identity. Original recordings are private and excluded. The 31.97 MB database export was checked for common identity and credential fields, preserved with a content hash and kept outside the submitted appendices because landmark tapes can be identifying movement data. Only curated interpretations and approved anonymized derivatives are used.")
add_body(doc, "Self-evaluation creates confirmation risk because the researcher has an interest in the artifact's success and supplied the expert rules. The mitigation is transparency rather than a false claim of independence: the report states the dual role, freezes permitted claims, reports contradictory evidence, avoids retrospective ratings and identifies second-expert review as future work.")

add_chapter(doc, 4, "Results and Discussion")
add_heading(doc, "4.1 ACP-STGAT Offline Results", 2)
acp_rows = [
    ["Normalized MPJPE / ADE", "0.07839 ± 0.00155", "0.09629", "0.10004"],
    ["Final displacement error", "0.13639 ± 0.00237", "0.16470", "0.22710"],
    ["Bone-length MAE", "0.00576 ± 0.00019", "approximately 0", "0.02487"],
    ["Model-only CPU latency", "median 9.88 ms; p95 12.81 ms", "-", "-"],
]
add_table(doc, "Table 4.1 ACP-STGAT model results and baselines", ["Metric", "ACP-STGAT mean ± SD", "Last pose", "Constant velocity"], acp_rows, [1.5, 1.8, 1.35, 1.45], "Coordinates are normalized. MPJPE and ADE are equivalent in this implementation; they are not separate findings.")
add_figure(doc, ROOT / "research/figures/verified/20260801/F1_acp_error_by_horizon.png", "Figure 4.1 ACP-STGAT normalized prediction error across the 30-frame forecast horizon on the participant-held-out Human3.6M-17 protocol. Error is not expressed in millimetres or as an accuracy percentage.", 5.8)
add_body(doc, "ACP-STGAT reduced mean prediction error by 18.6% relative to last pose and 21.6% relative to constant velocity. Final-frame error was 17.2% lower than last pose and 39.9% lower than constant velocity. Error increased with horizon for all approaches. Constant velocity was competitive in the first frames but accumulated greater drift, while the learned model better controlled the longer horizon. The across-seed standard deviations were small enough to support preliminary repeatability across the three runs.")
add_body(doc, "Robustness testing showed sensitivity to coordinate noise: noise with standard deviation 0.005 increased normalized MPJPE by about 18.4%, and standard deviation 0.010 increased it by about 50.4% relative to the clean seed-42 result. Simulated missing-landmark rates of 5% and 10% produced much smaller increases under the notebook's imputation procedure. This pattern supports confidence filtering and smoothing, but it cannot be generalized beyond the tested corruption model.")
add_body(doc, "The ONNX export passed parity with a maximum absolute PyTorch-ONNX difference of approximately 1.19 × 10^-7. Median model inference was 9.88 ms and p95 was 12.81 ms. These values are promising for deployment, but camera-to-feedback latency was not measured. The evidence supports a successful prototype benchmark result, not state-of-the-art performance or validated jab anticipation.")

add_heading(doc, "4.2 Temporal Phase-Classifier Results", 2)
phase_rows = [
    ["Accuracy", "87.84% ± 1.31 percentage points", "33.49%"],
    ["Balanced accuracy", "90.84% ± 0.71 percentage points", "14.29%"],
    ["Macro precision", "0.8803 ± 0.0113", "0.0478"],
    ["Macro recall", "0.9084 ± 0.0071", "0.1429"],
    ["Macro F1", "0.8915 ± 0.0093", "0.0717"],
    ["Weighted F1", "0.8786 ± 0.0133", "0.1680"],
]
add_table(doc, "Table 4.2 Temporal phase-classifier generated-bootstrap results", ["Metric", "ST-GCN/TCN mean ± SD", "Majority baseline"], phase_rows, [1.7, 2.7, 1.7], "All results are measured on generated bootstrap sessions and do not represent real-human accuracy.")
add_figure(doc, ROOT / "research/figures/verified/20260801/F2_phase_bootstrap_confusion_matrix.png", "Figure 4.2 Temporal phase-classifier confusion matrix for the generated-bootstrap evaluation set. The matrix verifies generator-defined pipeline behaviour only.", 5.3)
add_body(doc, "The learned classifier improved accuracy by 54.35 percentage points over the technique-majority baseline. Seed results were 87.42%, 86.79% and 89.31% accuracy. Seed 43 was selected for deployment because it had the highest validation macro F1 (0.9246), even though seed 44 had the highest test score; this avoids selecting on the test set.")
add_body(doc, "For the selected model, ordinary-phase F1 ranged from 0.842 for execution to 0.899 for retraction. Preparation recall was 0.793, while execution recall was 0.960 with lower precision of 0.750, indicating that execution was sometimes over-predicted. Tracking-lost achieved a nominal F1 of 1.0 but had support of only two test frames and is not emphasized. Jab technique-level accuracy was 83.99%, lower than the front-kick result of 89.84%, reinforcing the need for real grouped jab data.")
add_body(doc, "Synthetic boundary precision was 1.000, recall 0.952 and F1 0.974, with mean timing error of 0.954 frames. Coarse repetition recall was 0.8125 and repetition-count MAE was 0.375 per session; two multi-repetition jab sessions were under-counted. Ten-percent missing landmarks reduced macro F1 from 0.8867 to 0.8514. The ONNX test passed, with median model-only latency 5.74 ms and p95 30.49 ms. Each result remains generator-specific.")

add_heading(doc, "4.3 Automated Software Verification", 2)
software_rows = [
    ["Frontend", "23 test files", "129", "129", "Passed"],
    ["Backend", "Current suite", "24", "24", "Passed"],
    ["ACP ONNX", "Numerical parity", "Tolerance 1 × 10^-4", "Max diff 1.19 × 10^-7", "Passed"],
    ["Phase ONNX", "Deployment check", "Parity and inference", "Passed", "Passed"],
]
add_table(doc, "Table 4.3 Automated software verification", ["Area", "Scope", "Expected", "Observed", "Status"], software_rows, [1.0, 1.45, 1.35, 1.55, 0.75])
add_body(doc, "The tests verify implementation paths, data contracts and deterministic behaviour exercised by the suites. They provide strong evidence that the checked software behaves as specified under test conditions. They do not measure landmark correctness, coaching effectiveness or user outcomes. This distinction is important because a technically correct rule can still encode an unvalidated threshold, and a passing component can still fail when integrated with live camera timing.")

add_heading(doc, "4.4 P001 Expert Feasibility/Self-Evaluation Case", 2)
p001_rows = [
    ["Participant", "P001: researcher, developer and martial-arts expert"],
    ["Technique / device", "Jab / laptop camera"],
    ["Confirmed session", "practice-42, 1 August 2026"],
    ["Capture", "250 frames; 8.3 seconds"],
    ["Canonical repetitions", "3 completed"],
    ["Tracking quality", "96.0% recorded application measure"],
    ["Post-session clusters", "3 confirmed by P001 as good/correct"],
    ["Decoder changes", "22 frames changed from live labels"],
]
add_table(doc, "Table 4.4 P001 practice-42 evidence summary", ["Item", "Verified value"], p001_rows, [1.9, 4.2], "This is expert self-validation, not independent ground truth. Application scores are not accuracy.")
add_figure(doc, ROOT / "research/figures/verified/20260801/F4_p001_practice42_three_cluster_timeline.png", "Figure 4.3 Post-session three-cluster timeline for the confirmed P001 practice-42 jab case. P001 is the researcher/developer/martial-arts expert; confirmation is self-validation.", 5.9)
add_body(doc, "The wider screenshot set also documents the implemented interaction path across mode selection, dashboard review, Easy and Hard coaching, voice-mediated transitions, practice-set execution, post-session analysis and administrator diagnostics. Representative anonymized frames are preserved in Appendix A4. They establish that these interface states and outputs were displayed during the P001 feasibility case; they do not independently validate the numerical scores, model accuracy, coaching effectiveness or generalizability.")
add_body(doc, "The case demonstrates an end-to-end path from camera tracking through live interaction, session recording and post-session analysis. P001 confirmed the three displayed movement clusters as correct/good representations of the demonstrated repetitions. The analysis panel exposed the movement timeline, session accuracy trace, selected moments, phase labels, skeleton view and hand/face diagnostics. Train Easy and Train Hard screenshots additionally showed that the interface waited for responses, accepted microphone commands, prioritized corrections, and started a new session when requested.")
add_body(doc, "These observations support functional feasibility and qualitative alignment with the researcher's expert interpretation. They do not provide independent annotation, a usability scale or a controlled comparison. The 96% tracking field is an application diagnostic and not ground-truth landmark accuracy. The historical export contained many development and outlier sessions, which were excluded from the headline case rather than averaged into a misleading score.")

add_heading(doc, "4.5 Retained Integration Failure", 2)
failure_rows = [
    ["Canonical session summary", "3 completed repetitions", "Used for confirmed case"],
    ["Saved rule-engine summary", "0 completed repetitions", "Contradicts canonical summary"],
    ["Database repetition table", "1 row with implausible duration", "Persistence inconsistency"],
]
add_table(doc, "Table 4.5 Retained practice-42 data-layer inconsistency", ["Representation", "Observed value", "Interpretation"], failure_rows, [1.8, 2.05, 2.25])
add_body(doc, "The discrepancy is not removed from the evidence. It shows that the application can present a coherent canonical session analysis while other persistence layers disagree. This is a meaningful framework result because reliable cognitive-state and learner-memory claims depend on consistent state across components. The failure motivates a single versioned repetition contract, end-to-end integration tests and reconciliation checks at export time.")

add_heading(doc, "4.6 Evidence-to-Claim Reconciliation", 2)
claim_rows = [
    ["Supported with boundary", "4", "Baselines/horizon metrics; monocular uncertainty; design-science case; STGAT exclusion"],
    ["Partially supported", "5", "Perception accuracy; real-jab ACP; real phase boundaries; comparative decision benefit; threshold validation"],
    ["Unsupported", "0", "No controlled claim lacks all relevant evidence"],
    ["Closed evaluation actions", "4", "Baselines; retained inconsistency; LLM non-applicability; video exclusion"],
    ["Open but bounded", "8", "Explicit limitations/future evaluation; no additional data required for current feasibility conclusion"],
]
add_table(doc, "Table 4.6 Controlled claim readiness", ["Classification", "Count", "Interpretation"], claim_rows, [1.55, 0.65, 3.9])
add_body(doc, "The reconciliation prevents literature from substituting for local evidence. Laptop-camera perception is technically feasible and visibly implemented, but independent landmark and jab-form accuracy are absent. ACP-STGAT is a successful benchmark prototype, but real jab prediction and component ablation remain open. The phase pipeline has complete generated-data metrics, but no manually annotated real boundaries. Situation-awareness functions are implemented as computational gates, but comparative benefit has not been tested. Jab variables are supported by literature and expert traces, while threshold values remain researcher-authored.")

add_heading(doc, "4.7 Discussion Against the Research Question", 2)
add_body(doc, "The evidence answers the research question at the architecture and feasibility level. A hybrid framework can simulate selected computational functions by assigning different roles to perception, learned temporal evidence, deterministic constraints and memory. Learned models estimate phase and future pose; deterministic layers judge trust, enforce ordering, calculate form evidence and select an actionable response. Multi-timescale context prevents the system from treating every frame as independent, and structured storage supports post-session interpretation.")
add_body(doc, "The strongest evidence concerns implemented integration, reproducible offline model behaviour and software verification. The weakest evidence concerns external validity: one self-evaluating expert, one technique, no calibrated 3D ground truth, no condition-labelled rule-only/hybrid comparison, no end-to-end latency observations and no structured usability instrument. The scientifically appropriate conclusion is therefore feasibility, not effectiveness.")
add_body(doc, "Endsley's framework helps explain the information flow, but the implementation should be described as computational situation awareness rather than awareness in the human sense. Design-science methodology supports evaluating utility and integrity of the artifact while retaining negative evidence. Jab biomechanics sources justify multi-joint context but do not validate universal degree thresholds. Taken together, the literature and local evidence support the bounded claim that Combat Cognition is a functioning core architecture toward a broader research vision.")

add_heading(doc, "4.8 Limitations", 2)
for item in [
    "P001 is the researcher, developer and martial-arts expert; no independent participant or second expert was included.",
    "Only a jab and one laptop-camera setting were evaluated; results cannot be generalized to other techniques, people, viewpoints or devices.",
    "ACP-STGAT was evaluated on Human3.6M-17 rather than grouped real jab ground truth, and no ablation or advanced published-model comparison was completed.",
    "Temporal phase results use generated bootstrap sessions; real annotated phase and boundary accuracy is unknown.",
    "Only model-level ONNX latency was measured; full camera-to-feedback latency and sustained frame rate were not recorded.",
    "Historical records lack experimental-condition labels, preventing a controlled rule-only versus hybrid comparison.",
    "Expert-authored biomechanical thresholds have not been independently triangulated and cannot support clinical, force, energy or injury claims.",
    "No structured usability score was captured; qualitative commentary was not converted into retrospective ratings.",
    "The operational coach is deterministic; no current LLM performance can be claimed.",
    "The practice-42 repetition inconsistency shows unresolved integration risk across stored representations.",
]: add_bullet(doc, item)

add_chapter(doc, 5, "Conclusions and Recommendations")
add_heading(doc, "5.1 Conclusion", 2)
add_body(doc, "This study designed, implemented and evaluated Combat Cognition as a hybrid computational framework toward martial-artist cognitive simulation. The framework connects laptop-camera perception, expert-authored biomechanics, temporal phase estimation, future-pose prediction, four levels of contextual state, bounded situation-awareness gates, deterministic coaching and session memory. The result is an implemented architecture capable of interpreting and responding to a jab practice session while retaining the provenance and limits of its evidence.")
add_body(doc, "ACP-STGAT produced a promising participant-held-out benchmark result and outperformed last-pose and constant-velocity baselines in normalized coordinates. The temporal phase classifier substantially outperformed a majority baseline in generated bootstrap sessions and completed an ONNX deployment path. Automated tests verified the checked frontend and backend behaviour. The P001 practice-42 case demonstrated end-to-end functional use and expert-confirmed movement clustering, while simultaneously revealing a persistence inconsistency that was preserved as failure evidence.")
add_body(doc, "The research question is answered conditionally: a hybrid perception-awareness-decision architecture can simulate selected computational functions by combining learned temporal evidence with deterministic constraints, confidence gates, multi-timescale memory and explainable feedback selection. The work does not establish complete martial-artist cognition, coaching effectiveness, calibrated biomechanics, population usability or generalization. Combat Cognition is therefore best understood as an evidence-bounded core prototype and research platform.")

add_heading(doc, "5.2 Contributions", 2)
for item in [
    "An implemented end-to-end architecture linking perception, temporal prediction, ordered phase reasoning, multi-level context, situation-aware evidence selection, coaching and analytics.",
    "A reproducible ACP-STGAT benchmark evaluation with simple baselines, repeated seeds, robustness, ONNX parity and model-only latency.",
    "A complete generated-bootstrap temporal phase pipeline evaluation with class, boundary, repetition, robustness and deployment evidence.",
    "A deterministic structured coaching boundary that makes the operational reasoning mechanism explicit and keeps a future LLM optional.",
    "A curated P001 expert feasibility/self-evaluation case with confirmed clusters and a retained integration failure.",
    "A claim-control framework that separates literature, model, software and case evidence and prohibits a combined overall accuracy.",
]: add_bullet(doc, item)

add_heading(doc, "5.3 Recommendations", 2)
recommendations = [
    "Create a manually annotated, grouped real-jab dataset and freeze participant/session splits before windowing.",
    "Evaluate ACP-STGAT on jab-domain future frames, run component ablations and compare with established motion-prediction models under one protocol.",
    "Evaluate phase boundaries, edit score, segmental F1, repetition recovery and tracking-loss behaviour on real sequences with independent labels.",
    "Instrument camera capture, pose inference, model inference, decision and feedback timestamps to measure end-to-end latency and frame rate.",
    "Add explicit rule-only and hybrid condition labels and run a controlled replay or prospective comparison.",
    "Triangulate form variables and thresholds with independent martial-arts experts and, where physical claims are required, calibrated IMUs, force or motion-capture sensors.",
    "Repair repetition-state contracts across canonical, rule-engine and database layers and add export-time consistency tests.",
    "If an LLM is introduced, retain the structured evidence packet, schema validation, deterministic fallback, safety gates and blinded comparison against the current template baseline.",
    "Extend evaluation gradually to additional participants, viewpoints, devices and techniques without treating the present jab case as universal evidence.",
]
for item in recommendations: add_number(doc, item)

add_heading(doc, "5.4 Future Research Vision", 2)
add_body(doc, "The longer-term vision is a progressively richer computational model of martial-arts perception, anticipation, tactical choice, adaptation and personalized learning. Future sensor fusion could investigate movement efficiency through objective kinematics, force, impulse, muscle activation and physiological cost while preserving balance, recovery and practical effectiveness. Such work must compare defined technique variants under controlled conditions rather than declaring traditional movement inefficient from camera observations alone.")
add_body(doc, "A future local or domain-specific language model may improve explanation flexibility, but it should remain downstream of verified movement evidence and deterministic safety constraints. The principal research opportunity is not to replace every rule with a generative model; it is to learn which combination of perception, prediction, expert knowledge, uncertainty and memory produces reliable, interpretable coaching decisions.")

add_chapter(doc, "", "References")
# Replace the visually awkward blank chapter number with a normal heading.
last = doc.paragraphs[-1]
last.runs[0].text = "REFERENCES"
lit_path = ROOT / "research/outputs/literature/20260801/verified_literature_matrix.csv"
with lit_path.open(encoding="utf-8-sig", newline="") as f:
    refs = list(csv.DictReader(f))
for r in refs:
    authors = r["authors"]
    year = r["year"]
    title = r["title"]
    venue = r["venue"]
    locator = r["doi"] or r["primary_url"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(f"{authors} ({year}) '{title}', {venue}. {locator}"), 10)

add_chapter(doc, "", "Appendices")
doc.paragraphs[-1].runs[0].text = "APPENDICES"
appendix_sections = [
    ("Appendix A1: Architecture and Implementation Evidence", [
        "The architecture register records perception, biomechanics, L1-L4, situation-awareness, context, reasoning and analytics components with their code evidence, model/service, implementation status, limitation and future work. The full machine-readable register is retained in the research workspace.",
        "Key control: implemented does not mean empirically validated. Planned components, including an external or local LLM, remain classified separately from operational components.",
    ]),
    ("Appendix A2: ACP-STGAT Reproducibility Record", [
        "Frozen run: 20260731T061529Z. Seeds: 42, 43 and 44. Evaluation: participant-held-out Human3.6M-17, up to 10,000 test windows, 30-frame forecast. Metrics: normalized MPJPE/ADE, FDE and bone-length MAE. Baselines: last pose and constant velocity. Deployment: ONNX parity and CPU model-only latency.",
        "Large checkpoints and binary model files are retained in the local research package and are not embedded in this report.",
    ]),
    ("Appendix A3: Phase-Classifier Reproducibility Record", [
        "Frozen run: 20260731T091140Z. Generated bootstrap: 48 sessions (24 jab and 24 front kick). Seeds: 42, 43 and 44. Reported evidence includes classification, per-phase, boundary, repetition, robustness and ONNX deployment measures.",
        "The generated-data status is mandatory. No appendix item converts these values into real-human accuracy.",
    ]),
    ("Appendix A4: System Interface and Functional Evidence", [
        "The publication-ready supplementary appendix contains 51 distinct, account-name-masked screenshots retained after review of all 78 supplied images. They document visible controls, state transitions and diagnostic outputs in the P001 jab-only laptop-camera case. They are observational implementation evidence, not independent measurements of accuracy, effectiveness, causal reasoning or human-equivalent cognition. To avoid duplicating the same screenshots inside the report, the curated images are provided through the repository link below.",
    ]),
    ("Appendix A5: P001 Protocol and Curated Evidence", [
        "P001 is the researcher/developer/martial-arts expert. The case uses a jab and laptop camera. The selected practice-42 record is the confirmed latest/good case. Anonymized representative screenshots are included where they provide distinct functional evidence. Original recordings and repetitive or unnecessary screenshots are excluded.",
    ]),
    ("Appendix A6: Software Verification", [
        "The final evidence consolidation records 129/129 frontend assertions across 23 test files and 24/24 current backend tests. Verification covers checked code paths and contracts; it does not establish live coaching correctness.",
    ]),
    ("Appendix A7: Database Export and Failure Evidence", [
        "The verified export contained 42 practice sessions, 58 recorded repetitions, 22 landmark tapes with 7,277 frames, 185 training sessions, 140 training-step attempts and 3,623 feedback events. These records include development history and are not aggregated into a controlled accuracy result.",
        "The 31.97 MB raw JSON and potentially identifying movement tapes are excluded. The interpretation and practice-42 inconsistency remain part of the audit trail.",
    ]),
    ("Appendix A8: Literature and Claim Audit", [
        "The verified literature matrix contains 22 sources: 21 included or included with caution and one direct-support exclusion. Thirty-nine page/section-specific extracts connect source propositions, limitations and local evidence requirements.",
    ]),
    ("Appendix A9: Evaluation Actions and Evidence Locks", [
        "Twelve evaluation actions were audited. Four were closed or complete, and eight remained open but bounded as limitations/future evaluation. No new participant, recording or retrospective score was invented for the present feasibility conclusion.",
    ]),
]
for heading, paras in appendix_sections:
    add_heading(doc, heading, 2)
    for para in paras: add_body(doc, para)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    appendix_id = heading.split(":", 1)[0].replace("Appendix ", "")
    set_font(p.add_run(f"Supplementary evidence ({appendix_id}): "), 12, bold=True)
    appendix_url = f"{APPENDIX_BASE_URL}/{APPENDIX_FOLDERS[heading]}"
    add_hyperlink(p, f"Open Appendix {appendix_id} evidence folder on GitHub", appendix_url)

# Update fields on open and remove metadata traces.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "false")
doc.core_properties.title = TITLE.title()
doc.core_properties.subject = "DS 5299 - Independent Study"
doc.core_properties.author = "I.T.M.S.S.B. Thennakoon"
doc.core_properties.keywords = "Combat Cognition, martial arts, pose estimation, temporal reasoning, situation awareness"

# Keep headings/captions with following content and avoid widows where possible.
for p in doc.paragraphs:
    pPr = p._p.get_or_add_pPr()
    widow = pPr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        pPr.append(widow)
    widow.set(qn("w:val"), "1")

doc.save(DOCX)
print(DOCX)
