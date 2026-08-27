from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor


OUT = Path(__file__).resolve().parent
DOCX = OUT / "Combat_Cognition_Thesis_I.T.M.S.S.B.Thennakoon.docx"
BASE_URL = "https://github.com/SachithBandaraThennakoon/martial-art-ai/tree/main/research/Appendix"

APPENDICES = {
    "Appendix A1: Architecture and Implementation Evidence": "A1-Architecture-and-Implementation",
    "Appendix A2: ACP-STGAT Reproducibility Record": "A2-ACP-STGAT-Reproducibility",
    "Appendix A3: Phase-Classifier Reproducibility Record": "A3-Phase-Classifier-Reproducibility",
    "Appendix A4: System Interface and Functional Evidence": "A4-System-Interface-and-Functional-Evidence",
    "Appendix A5: P001 Protocol and Curated Evidence": "A5-P001-Protocol-and-Curated-Evidence",
    "Appendix A6: Software Verification": "A6-Software-Verification",
    "Appendix A7: Database Export and Failure Evidence": "A7-Database-Export-and-Failure-Evidence",
    "Appendix A8: Literature and Claim Audit": "A8-Literature-and-Claim-Audit",
    "Appendix A9: Evaluation Actions and Evidence Locks": "A9-Evaluation-Actions-and-Evidence-Locks",
}


def set_font(run, *, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    properties.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document(DOCX)

# Remove the former single root-level link and any links from an earlier run.
for paragraph in list(doc.paragraphs):
    if paragraph.text.strip().startswith("Supplementary evidence (A") or paragraph.text.strip().startswith("Supplementary appendices:"):
        paragraph._element.getparent().remove(paragraph._element)

for heading, folder in APPENDICES.items():
    paragraphs = doc.paragraphs
    heading_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == heading)
    next_heading = next(
        (paragraph for paragraph in paragraphs[heading_index + 1:] if paragraph.text.strip() in APPENDICES),
        None,
    )

    link_paragraph = doc.add_paragraph()
    link_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    link_paragraph.paragraph_format.space_after = Pt(6)
    appendix_id = heading.split(":", 1)[0].replace("Appendix ", "")
    set_font(link_paragraph.add_run(f"Supplementary evidence ({appendix_id}): "), bold=True)
    url = f"{BASE_URL}/{folder}"
    add_hyperlink(link_paragraph, url, url)

    if next_heading is not None:
        next_heading._p.addprevious(link_paragraph._p)

# Remove obsolete hyperlink relationships left by the former root link or a rerun.
used_relationship_ids = set(doc._element.xpath(".//w:hyperlink/@r:id"))
for relationship_id, relationship in list(doc.part.rels.items()):
    if relationship.reltype == RT.HYPERLINK and relationship_id not in used_relationship_ids:
        doc.part.drop_rel(relationship_id)

doc.save(DOCX)
print(f"Added {len(APPENDICES)} appendix links to {DOCX}")
