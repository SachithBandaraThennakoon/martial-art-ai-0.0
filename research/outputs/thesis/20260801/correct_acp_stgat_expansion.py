from pathlib import Path
from docx import Document


DOCX = Path(__file__).with_name("Combat_Cognition_Thesis_I.T.M.S.S.B.Thennakoon.docx")
OLD = "Anticipatory Cognitive Pose Spatial-Temporal Graph Attention model"
NEW = "Action-Conditioned Physics-Informed Spatio-Temporal Graph Attention Transformer"


def replace_in_paragraph(paragraph):
    full = "".join(run.text for run in paragraph.runs)
    if OLD not in full:
        return 0
    replacement = full.replace(OLD, NEW)
    if paragraph.runs:
        paragraph.runs[0].text = replacement
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replacement)
    return full.count(OLD)


document = Document(DOCX)
count = 0
for paragraph in document.paragraphs:
    count += replace_in_paragraph(paragraph)
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count += replace_in_paragraph(paragraph)

if count != 1:
    raise RuntimeError(f"Expected exactly one replacement, found {count}")

document.save(DOCX)
print(f"Updated {count} occurrence in {DOCX}")
